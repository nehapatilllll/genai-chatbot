# modules/user_docs.py
import uuid, os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from PyPDF2 import PdfReader 
from docx import Document as DocxDocument

EMBEDDING = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def ingest_user_file(file_obj, filename):
    """Parse PDF/DOCX/TXT and return (vectorstore, full_text)."""
    text = ""
    filename = filename.lower()
    if filename.endswith(".pdf"):
        text = "\n".join(p.extract_text() or "" for p in PdfReader(file_obj).pages)
    elif filename.endswith(".docx"):
        doc = DocxDocument(file_obj)
        text = "\n".join(p.text for p in doc.paragraphs)
    else:  # TXT
        text = file_obj.read().decode()

    if not text.strip():
        return None, ""
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = [Document(page_content=chunk) for chunk in splitter.split_text(text)]
    vectorstore = FAISS.from_documents(docs, EMBEDDING)
    return vectorstore, text