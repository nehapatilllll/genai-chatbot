# modules/backend_store.py
import os
import uuid
import networkx as nx
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings

# ------------------------------------------------------------------
# Config
# ------------------------------------------------------------------
BACKEND_INDEX_DIR = "faiss_backend"        # folder that will hold *.faiss
EMBEDDING = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# ------------------------------------------------------------------
# Vector store helpers
# ------------------------------------------------------------------
def get_backend_vectorstore():
    """
    Returns a FAISS vectorstore.
    Loads from disk if exists; otherwise returns a tiny empty store.
    """
    if os.path.isdir(BACKEND_INDEX_DIR):
        return FAISS.load_local(
            BACKEND_INDEX_DIR,
            EMBEDDING,
            allow_dangerous_deserialization=True
        )
    # Empty store for first run (will be overwritten on ingest)
    return FAISS.from_texts(["_init_"], EMBEDDING)

@st.cache_resource
def get_backend_graph():
    """Singleton NetworkX graph stored in Streamlit session."""
    if "backend_graph" not in st.session_state:
        st.session_state.backend_graph = nx.DiGraph()
    return st.session_state.backend_graph

# ------------------------------------------------------------------
# Ingestion
# ------------------------------------------------------------------
def ingest_into_backend(file_obj, filename):
    """
    Reads PDF / DOCX / TXT, chunks, embeds, and stores into FAISS + Graph.
    Returns number of chunks added.
    """
    text = ""
    filename = filename.lower()

    # ---- text extraction ----
    if filename.endswith(".pdf"):
        from PyPDF2 import PdfReader
        text = "\n".join(p.extract_text() or "" for p in PdfReader(file_obj).pages)
    elif filename.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(file_obj)
        text = "\n".join(p.text for p in doc.paragraphs)
    else:  # TXT
        text = file_obj.read().decode()

    if not text.strip():
        return 0

    # ---- chunk & embed ----
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = [Document(page_content=chunk,
                     metadata={"source": filename,
                               "chunk_id": str(uuid.uuid4())})
            for chunk in splitter.split_text(text)]

    vs = get_backend_vectorstore()
    vs.add_documents(docs)
    vs.save_local(BACKEND_INDEX_DIR)

    # ---- semantic graph ----
    G = get_backend_graph()
    file_node = f"file:{filename}"
    G.add_node(file_node, type="file", name=filename)
    for d in docs:
        chunk_node = d.metadata["chunk_id"]
        G.add_node(chunk_node, type="chunk", text=d.page_content[:100] + "...")
        G.add_edge(file_node, chunk_node)

    return len(docs)